import tkinter as tk
from tkinter import ttk, messagebox
import threading
import queue
import os
import pickle
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import json

def setup():
    """Загружает все необходимые модели, данные и API ключи."""
    print("Начало настройки системы...")

    load_dotenv()
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        print("ОШИБКА: API ключ GEMINI не найден. Создайте файл .env и добавьте GEMINI_API_KEY=ваш_ключ")
        return None, None, None, None
    genai.configure(api_key=gemini_api_key)
    print("API ключ Gemini загружен.")

    embedding_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
    print("Модель для эмбеддингов загружена.")
    
    try:
        faiss_index = faiss.read_index("faiss_index.bin")
        print(f"Векторная база FAISS загружена. В ней {faiss_index.ntotal} документов.")
    except Exception as e:
        print(f"ОШИБКА: не удалось загрузить faiss_index.bin. Убедитесь, что файл существует. {e}")
        return None, None, None, None
        
    try:
        with open("docs.pkl", "rb") as f:
            documents = pickle.load(f)
        print(f"Тексты документов (чанки) загружены. Всего {len(documents)} фрагментов.")
    except Exception as e:
        print(f"ОШИБКА: не удалось загрузить docs.pkl. Убедитесь, что файл существует. {e}")
        return None, None, None, None

    print("Настройка системы завершена.\n")
    return embedding_model, faiss_index, documents, genai.GenerativeModel("gemini-1.5-flash")

def search(query, k, embedding_model, faiss_index, documents):
    query_vector = embedding_model.encode([query])
    distances, indices = faiss_index.search(query_vector.astype('float32'), k)
    
    results = [documents[i] for i in indices[0]]
    return results

def get_context_for_phys_culture(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретная игра, упражнение или комплекс для детей {age_group} на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_phys_culture_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    master_prompt = f"""
ТЫ — ЭКСПЕРТ-МЕТОДИСТ, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры игр и методик из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Твой ответ должен быть сфокусирован на раскрытии "Ключевых тем для изучения". Также обязательно включи 1-2 активности для "Тем для закрепления".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" наиболее подходящие игры и упражнения для раскрытия каждой темы. Если материалы для какой-то темы не нашлись (контекст пуст), используй свой экспертный опыт, чтобы предложить подходящую активность.
3.  СТРУКТУРИРУЙ: Организуй ответ по тематическим блокам: "Основные движения", "Общеразвивающие упражнения", "Подвижная игра", "Спортивные упражнения".
4.  ДЕТАЛИЗАЦИЯ: Для каждого блока обязательно заполни подзаголовки: "Цели:", "Упражнения:", "Инвентарь:", "Ход игры:" (если применимо).
5.  СТИЛЬ: Текст должен быть четким, практичным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_speech_dev(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретная игра, упражнение или методика для детей {age_group} на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_speech_dev_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    master_prompt = f"""
ТЫ — ЭКСПЕРТ-МЕТОДИСТ, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры игр и методик из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Твой ответ должен быть сфокусирован на раскрытии "Ключевых тем для изучения". Также обязательно включи 1-2 активности для "Тем для закрепления".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" наиболее подходящие игры и упражнения для раскрытия каждой темы. Если материалы для какой-то темы не нашлись, используй свой экспертный опыт, чтобы предложить подходящую активность.
3.  СТРУКТУРИРУЙ: Организуй ответ по тематическим блокам: "Тематический словарь", "Звуковая культура речи", "Грамматический строй", "Связная речь".
4.  ДЕТАЛИЗАЦИЯ: Для каждого блока обязательно заполни подзаголовки: "Цели:", "Содержание работы:", "Материалы:".
5.  СТИЛЬ: Текст должен быть четким, практичным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_literature(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретное литературное произведение, сказка, стих или потешка для детей {age_group} на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_literature_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    example_activities_str = ", ".join(monthly_plan.get("example_activities", []))
    master_prompt = f"""
ТЫ — ЭКСПЕРТ-МЕТОДИСТ и детский литературовед, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы и жанры: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
- Рекомендуемые произведения (если есть в плане): {example_activities_str if example_activities_str else "Подобрать самостоятельно на основе тем"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Твой ответ должен быть сфокусирован на "Ключевых темах и жанрах". Предложи 2-3 произведения, соответствующие этим темам и/или "Рекомендуемым произведениям".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" информацию о том, как работать с выбранными произведениями. Если материалы не нашлись, используй свой экспертный опыт.
3.  СТРУКТУРА: Для каждого произведения создай отдельный блок.
4.  ДЕТАЛИЗАЦИЯ: Для каждого блока обязательно заполни подзаголовки: "Цели:", "Содержание работы:", "Материалы:".
5.  СТИЛЬ: Текст должен быть четким, практичным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_math(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретная дидактическая игра или упражнение для детей {age_group} на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_math_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    example_activities_str = ", ".join(monthly_plan.get("example_activities", []))
    master_prompt = f"""
ТЫ — ЭКСПЕРТ-МЕТОДИСТ, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
- Примеры рекомендуемых игр: {example_activities_str if example_activities_str else "Подобрать самостоятельно"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры игр и методик из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Твой ответ должен быть сфокусирован на раскрытии "Ключевых тем для изучения". Также обязательно включи 1-2 активности для "Тем для закрепления".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" и "Примеров рекомендуемых игр" наиболее подходящие дидактические игры для каждой темы.
3.  СТРУКТУРИРУЙ: Организуй ответ по тематическим блокам, соответствующим темам (например, "Количество и счет", "Геометрические фигуры", "Величина").
4.  ДЕТАЛИЗАЦИЯ: Для каждого блока обязательно заполни подзаголовки: "Цели:", "Содержание работы:", "Материалы:".
5.  СТИЛЬ: Текст должен быть четким, практичным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_art(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретное занятие, техника или поделка для детей {age_group} на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=1, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_art_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    example_activities_str = ", ".join(monthly_plan.get("example_activities", []))
    master_prompt = f"""
ТЫ — ЭКСПЕРТ-МЕТОДИСТ, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
- Примеры рекомендуемых активностей: {example_activities_str if example_activities_str else "Подобрать самостоятельно"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Твой ответ должен содержать 4 раздела: 1. Рисование, 2. Лепка, 3. Аппликация, 4. Конструирование. Для каждого раздела выбери одну из "Ключевых тем" или "Тем для закрепления".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" описание техник и хода работы для выбранных тем. Если материалы не нашлись, используй свой экспертный опыт.
3.  ДЕТАЛИЗАЦИЯ: Внутри каждого из четырех разделов обязательно используй подзаголовки: "Тема:", "Цели:", "Содержание работы:", "Материалы:", "Безопасность:".
4.  СТИЛЬ: Текст должен быть четким, практичным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_music(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретная песня, танец, музыкальная игра или упражнение для детей {age_group} на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_music_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    example_activities_str = ", ".join(monthly_plan.get("example_activities", []))
    master_prompt = f"""
ТЫ — ВЫСОКОКВАЛИФИЦИРОВАННЫЙ МУЗЫКАЛЬНЫЙ РУКОВОДИТЕЛЬ, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
- Примеры рекомендуемого репертуара: {example_activities_str if example_activities_str else "Подобрать самостоятельно"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Твой ответ должен быть сфокусирован на "Ключевых темах для изучения". Также обязательно включи 1-2 активности для "Тем для закрепления".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" и "Примеров репертуара" наиболее подходящие песни, пьесы и игры для каждой темы.
3.  СТРУКТУРИРУЙ: Организуй ответ по 4 разделам: 1. Слушание, 2. Пение, 3. Музыкально-ритмические движения, 4. Игра на инструментах.
4.  ДЕТАЛИЗАЦИЯ: Внутри каждого раздела обязательно используй подзаголовки: "Цели:", "Репертуар:", "Содержание работы:", "Материалы:".
5.  СТИЛЬ: Текст должен быть четким, практичным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_kazakh_lang(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретная лексическая тема, игра или упражнение для детей {age_group} по казахскому языку на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_kazakh_lang_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    master_prompt = f"""
ТЫ — ОПЫТНЫЙ ПРЕПОДАВАТЕЛЬ КАЗАХСКОГО ЯЗЫКА, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые лексические темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СТРУКТУРИРУЙ ПЛАН ПО ТЕМАМ: Раздели свой ответ на блоки по "Ключевым лексическим темам".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" конкретные слова для 'Сөздік минимум' и примеры игр для 'Жұмыс мазмұны'. Если материалы не нашлись, используй свой экспертный опыт.
3.  ДЕТАЛИЗАЦИЯ: Внутри каждого блока обязательно используй подзаголовки: "Мақсаттар (Цели):", "Сөздік минимум (Лексический минимум):", "Жұмыс мазмұны (Содержание работы):", "Материалдар (Материалы):".
4.  СТИЛЬ: Текст должен быть четким, практичным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_world(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретное занятие, беседа, наблюдение или дидактическая игра для детей {age_group} на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_world_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    master_prompt = f"""
ТЫ — ЭКСПЕРТ-МЕТОДИСТ, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Раздели свой ответ на 2-3 блока по основным "Ключевым темам для изучения".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" конкретные примеры наблюдений, бесед и дидактических игр для каждой темы. Если материалы не нашлись, используй свой экспертный опыт.
3.  ДЕТАЛИЗАЦИЯ: Внутри каждого блока обязательно используй подзаголовки: "Цели:", "Содержание работы:", "Материалы:".
4.  СТИЛЬ: Текст должен быть четким, познавательным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def get_context_for_literacy(embedding_model, faiss_index, documents, age_group, month, monthly_plan):
    print(f"Поиск методик по плану для: {age_group} / {month}")
    key_topics = monthly_plan.get("key_topics", [])
    reinforcement_topics = monthly_plan.get("reinforcement_topics", [])
    all_topics_for_search = key_topics + reinforcement_topics
    
    all_chunks = []
    if not all_topics_for_search:
        print("  - В карте на этот месяц нет тем для поиска. Контекст будет пустым.")
        return ""
        
    for topic in all_topics_for_search:
        query = f"Конкретная игра или упражнение для детей {age_group} по обучению грамоте на тему: '{topic}'"
        print(f"  - Поиск по теме: '{topic}'")
        results = search(query, k=2, embedding_model=embedding_model, faiss_index=faiss_index, documents=documents)
        for chunk in results:
            all_chunks.append(f"[Пример методики по теме '{topic}']: {chunk.page_content if hasattr(chunk, 'page_content') else chunk}")

    print(f"Найдено релевантных фрагментов: {len(all_chunks)}")
    context = "\n\n---\n\n".join(all_chunks)
    return context

def generate_literacy_cell_prompt(context, age_group, month, monthly_plan):
    print(f"Составление промпта для: {age_group} / {month}")
    key_topics_str = ", ".join(monthly_plan.get("key_topics", []))
    reinforcement_topics_str = ", ".join(monthly_plan.get("reinforcement_topics", []))
    master_prompt = f"""
ТЫ — ОПЫТНЫЙ ЛОГОПЕД-МЕТОДИСТ, который составляет план занятия СТРОГО ПО ЗАДАННОМУ УЧЕБНОМУ ПЛАНУ.
ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ЭТОТ МЕСЯЦ ({month}):
---
- Ключевые темы для изучения: {key_topics_str}
- Темы для закрепления: {reinforcement_topics_str if reinforcement_topics_str else "Нет"}
---
ОПОРНЫЕ МАТЕРИАЛЫ (Примеры из базы знаний, найденные по темам из ТЗ):
---
{context}
---
ИНСТРУКЦИИ ПО ГЕНЕРАЦИИ:
1.  СЛЕДУЙ ПЛАНУ: Твой ответ должен быть сфокусирован на раскрытии "Ключевых тем для изучения". Также обязательно включи 1-2 активности для "Тем для закрепления".
2.  ИСПОЛЬЗУЙ МАТЕРИАЛЫ: Возьми из "ОПОРНЫХ МАТЕРИАЛОВ" конкретные примеры игр и упражнений для каждой темы.
3.  СТРУКТУРА: Организуй ответ в виде единого текста, но логически сгруппируй активности по темам.
4.  ДЕТАЛИЗАЦИЯ: Обязательно используй подзаголовки: "Цели:", "Содержание работы:", "Материалы:".
5.  СТИЛЬ: Текст должен быть четким, методически верным, без Markdown-форматирования и лишних пустых строк.
ПРЕДОСТАВЬ ГОТОВЫЙ, ДЕТАЛЬНЫЙ ТЕКСТ ДЛЯ ЯЧЕЙКИ, ВЫПОЛНЕННЫЙ ПО ТЕХНИЧЕСКОМУ ЗАДАНИЮ:
"""
    print("Промпт создан. Отправка в API.")
    return master_prompt

def create_document_header(doc, group_name, year):
    doc.add_paragraph('Согласовано').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f'Перспективный план организованной деятельности на {year} учебный год', style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Группа: {group_name}', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

def setup_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.layout_algorithm = 1
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Месяц'
    hdr_cells[1].text = 'Образовательная область'
    hdr_cells[2].text = 'Задачи организованной деятельности'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    widths = (Cm(2.5), Cm(3.5), Cm(10.0)) 
    for i, width in enumerate(widths):
        table.columns[i].width = width
        for cell in table.columns[i].cells:
            cell.width = width
            
    return table

def clean_text(text):
    text = text.replace('**', '')
    lines = text.split('\n')
    non_empty_lines = [line for line in lines if line.strip() != '']
    final_text = '\n'.join(non_empty_lines)
    return final_text

def add_row_to_table(table, month, area, content, is_first_entry_for_month=False):
    row_cells = table.add_row().cells
    row_cells[1].text = area
    row_cells[2].text = content
    
    month_cell = row_cells[0]
    if is_first_entry_for_month:
        month_cell.text = month
    else:
        month_cell.merge(table.cell(len(table.rows) - 2, 0)) 
    
    print(f"Ячейка добавлена в документ: {month} / {area}")

def run_generation_process(age_group, update_queue):
    """Эта функция содержит всю логику из старого __main__ и "общается" с GUI через очередь."""
    try:
        update_queue.put(("status", "Шаг 0/4: Настройка системы..."))
        embedding_model, faiss_index, documents, generative_model = setup()
        if not all((embedding_model, faiss_index, documents, generative_model)):
            raise Exception("Ошибка инициализации моделей или базы знаний.")
        
        update_queue.put(("status", "Шаг 1/4: Загрузка учебной программы..."))
        with open("curriculum_map.json", "r", encoding="utf-8") as f:
            curriculum_map = json.load(f)
        
        YEAR = "2025-2026"
        ALL_MONTHS = ["Сентябрь", "Октябрь", "Ноябрь", "Декабрь", "Январь", "Февраль", "Март", "Апрель", "Май"]
        plan_for_age_group = curriculum_map.get(age_group)
        if not plan_for_age_group:
            raise Exception(f"Не найдена программа для группы '{age_group}'")

        total_steps = sum(1 for area, plans in plan_for_age_group.items() for p in plans)
        steps_completed = 0

        FUNCTION_MAP = {
            "Физическая культура": "phys_culture",
            "Развитие речи": "speech_dev",
            "Художественная литература": "literature",
            "Основы грамоты": "literacy",
            "Основы математики": "math",
            "Рисование/Лепка/Аппликация/Конструирование": "art",
            "Музыка": "music",
            "Казахский язык": "kazakh_lang",
            "Ознакомление с окружающим миром": "world"
        }

        update_queue.put(("status", "Шаг 2/4: Создание Word документа..."))
        document = Document()
        create_document_header(document, age_group, YEAR)
        plan_table = setup_table(document)
        
        update_queue.put(("status", "Шаг 3/4: Начало генерации контента..."))
        
        for month in ALL_MONTHS:
            is_first_entry_for_month = True
            for area, monthly_plans in plan_for_age_group.items():
                monthly_plan = next((p for p in monthly_plans if p['month'] == month), None)
                if not monthly_plan: continue
                
                status_msg = f"Генерация: {month} / {area}"
                update_queue.put(("status", status_msg))
                
                func_name_suffix = FUNCTION_MAP.get(area)
                get_context_func = globals().get(f"get_context_for_{func_name_suffix}")
                get_prompt_func = globals().get(f"generate_{func_name_suffix}_cell_prompt")

                if get_context_func and get_prompt_func:
                    context = get_context_func(embedding_model, faiss_index, documents, age_group=age_group, month=month, monthly_plan=monthly_plan)
                    prompt = get_prompt_func(context, age_group=age_group, month=month, monthly_plan=monthly_plan)
                    
                    response = generative_model.generate_content(prompt)
                    raw_content = response.text
                    cell_content = clean_text(raw_content)
                    
                    add_row_to_table(plan_table, month, area, cell_content, is_first_entry_for_month=is_first_entry_for_month)
                    is_first_entry_for_month = False
                
                steps_completed += 1
                progress = (steps_completed / total_steps) * 100
                update_queue.put(("progress", progress))

        update_queue.put(("status", "Шаг 4/4: Сохранение файла..."))
        safe_age_group = age_group.replace(' ', '_').replace('(', '').replace(')', '').replace('/', '_')
        output_filename = f"Годовой_Перспективный_план_{safe_age_group}.docx"
        document.save(output_filename)
        
        update_queue.put(("status", f"Готово! План сохранен: {output_filename}"))
        update_queue.put(("done", output_filename))

    except Exception as e:
        update_queue.put(("error", str(e)))

class PlanGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор Перспективных Планов")
        self.root.geometry("500x250")
        
        self.update_queue = queue.Queue()

        self.age_groups = [
            "Младшая группа (2-3 года)",
            "Средняя группа (3-4 года)",
            "Старшая группа (4-5 лет)",
            "Предшкольная группа (5-6 лет)"
        ]
        self._create_widgets()

    def _create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        ttk.Label(main_frame, text="Выберите возрастную группу:").pack(pady=(0, 5))
        self.age_combo = ttk.Combobox(main_frame, values=self.age_groups, state="readonly")
        self.age_combo.pack(fill=tk.X, pady=(0, 10))
        self.age_combo.set(self.age_groups[2])
        self.start_button = ttk.Button(main_frame, text="Начать генерацию", command=self.start_generation)
        self.start_button.pack(pady=10)
        self.progress_bar = ttk.Progressbar(main_frame, orient="horizontal", length=300, mode="determinate")
        self.progress_bar.pack(pady=5)
        self.status_label = ttk.Label(main_frame, text="Готов к работе.")
        self.status_label.pack(pady=5)

    def start_generation(self):
        selected_group = self.age_combo.get()
        if not selected_group:
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите возрастную группу.")
            return

        self.start_button.config(state="disabled")
        self.progress_bar["value"] = 0
        self.status_label.config(text="Инициализация...")
        
        self.generation_thread = threading.Thread(target=run_generation_process, args=(selected_group, self.update_queue))
        self.generation_thread.start()
        
        self.root.after(100, self.check_queue)

    def check_queue(self):
    
        while not self.update_queue.empty():
            try:
                message = self.update_queue.get_nowait()
                msg_type, msg_data = message
                
                if msg_type == "progress":
                    self.progress_bar["value"] = msg_data
                elif msg_type == "status":
                    self.status_label.config(text=msg_data)
                elif msg_type == "done":
                    
                    self.progress_bar["value"] = 100 
                    self.start_button.config(state="normal")
                    messagebox.showinfo("Успех", f"Генерация успешно завершена!\nФайл сохранен как: {msg_data}")
                    return
                elif msg_type == "error":
                    self.start_button.config(state="normal")
                    messagebox.showerror("Ошибка", f"Произошла ошибка:\n{msg_data}")
                    return 

            except queue.Empty:
                pass

        if self.generation_thread.is_alive():
            self.root.after(100, self.check_queue)
        else:
            if self.start_button['state'] == 'disabled':
                self.start_button.config(state="normal")
                self.status_label.config(text="Процесс завершен.")

if __name__ == "__main__":
    root = tk.Tk()
    app = PlanGeneratorApp(root)
    root.mainloop()

